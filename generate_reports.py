import os
import docx
from docx.shared import Pt


def report_generator(cleaned_data,output_folder):
    """
    Generates individual student academic reports in Word format.
    """

    report_dir=os.path.join(output_folder, "reports")
    individual_dir = os.path.join(report_dir, "individual")

    os.makedirs(individual_dir, exist_ok=True)

    for row in cleaned_data:
        student_name = row["full_name"]
        safe_name = student_name.replace(" ", "_")
        student_id = row["student_id"]

        filename = f"{student_id}_{safe_name}_Report.docx"
        filepath = os.path.join(individual_dir, filename)

        document = docx.Document()

        # ----- Title -----
        title = document.add_paragraph("Student Academic Report")
        title_format = title.runs[0].font
        title_format.size = Pt(16)
        title_format.bold = True
        title.alignment = 1  # Center alignment

        document.add_paragraph("")  # Spacer

        # ----- Student Information Section -----
        header = document.add_paragraph("Student Information")
        header.runs[0].bold = True

        document.add_paragraph(f"Full Name: {student_name}")
        document.add_paragraph(f"Student ID: {student_id}")
        document.add_paragraph(f"Course: {row['course']}")

        document.add_paragraph("")  # Spacer

        # ----- Scores Section -----
        header = document.add_paragraph("Assessment Scores")
        header.runs[0].bold = True

        document.add_paragraph(f"Score 1: {row['score1']}")
        document.add_paragraph(f"Score 2: {row['score2']}")
        document.add_paragraph(f"Score 3: {row['score3']}")

        document.add_paragraph("")  # Spacer

        # ----- Summary Section -----
        header = document.add_paragraph("Performance Summary")
        header.runs[0].bold = True

        document.add_paragraph(f"Total Score: {row['total_score']}")
        document.add_paragraph(f"Average Score: {row['average_score']}")
        document.add_paragraph(f"Final Grade: {row['grade']}")

        document.save(filepath)
